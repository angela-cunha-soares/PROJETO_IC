# -*- coding: utf-8 -*-
"""Etapa 7 do cronograma - geração do Relatório Final em Word (.docx).

Monta um relatório de resultados (~20 páginas) a partir das tabelas em
``reports/tables/`` e das figuras em ``reports/figures/`` produzidas pelas
etapas anteriores do pipeline. Usa python-docx (mesma stack Python do projeto).

Pré-requisitos: rodar antes, nesta ordem,
    run_preprocess.py -> run_outliers.py -> run_descritiva.py
    -> run_train.py -> run_evaluate.py

Saída:
    * ``reports/relatorio_final.docx``

Uso:
    python scripts/gerar_relatorio_docx.py
"""
from __future__ import annotations

import logging
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT / "src"))

import config  # noqa: E402

LOG = logging.getLogger("gerar_relatorio_docx")
FIG = config.FIGURES_DIR
TAB = config.TABLES_DIR
AZUL = RGBColor(0x1F, 0x4E, 0x79)
CINZA = RGBColor(0x59, 0x59, 0x59)

_fig_n = 0
_tab_n = 0


def h1(doc, texto):
    p = doc.add_heading(texto, level=1)
    for r in p.runs:
        r.font.color.rgb = AZUL
    return p


def h2(doc, texto):
    p = doc.add_heading(texto, level=2)
    for r in p.runs:
        r.font.color.rgb = AZUL
    return p


def para(doc, texto, *, justify=True, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.italic = italic
    r.font.size = Pt(size)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def figura(doc, arquivo, legenda, *, largura=6.1):
    global _fig_n
    caminho = FIG / arquivo
    if not caminho.is_file():
        LOG.warning("Figura ausente: %s (pulada)", arquivo)
        return
    _fig_n += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(caminho), width=Inches(largura))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figura {_fig_n}. {legenda}")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = CINZA


def _set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear",
                                         qn("w:color"): "auto",
                                         qn("w:fill"): hexcolor})
    tcPr.append(shd)


def tabela(doc, df, legenda, *, col_labels=None, fontsize=8.5, num=True):
    global _tab_n
    if num:
        _tab_n += 1
        cap = doc.add_paragraph()
        r = cap.add_run(f"Tabela {_tab_n}. {legenda}")
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = CINZA

    cols = list(df.columns)
    labels = col_labels or cols
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for j, lab in enumerate(labels):
        hdr[j].text = str(lab)
        for p in hdr[j].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(fontsize)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(hdr[j], "1F4E79")
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(cols):
            val = row[c]
            if isinstance(val, float):
                txt = f"{val:.3g}" if val == val else "-"
            else:
                txt = str(val)
            cells[j].text = txt
            for p in cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(fontsize)
    return t


def _read(nome):
    p = TAB / nome
    return pd.read_csv(p) if p.is_file() else None


def capa(doc):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Aplicação de Machine Learning para Análise da "
                  "Qualidade da Água na Bacia Hidrográfica do PCJ")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = AZUL

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Relatório Final de Iniciação Científica")
    r.font.size = Pt(15)
    r.font.color.rgb = CINZA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Estação de Tratamento de Água do rio Piracicaba (SEMAE) "
                  "— série histórica 2009–2024")
    r.italic = True
    r.font.size = Pt(11)

    for _ in range(8):
        doc.add_paragraph()
    for campo in ("Autor(a): ______________________________",
                  "Orientador(a): __________________________",
                  "Instituição / Programa: __________________",
                  "Parceria: Centro de Energia Nuclear na Agricultura (CENA)"):
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pp.add_run(campo)
        rr.font.size = Pt(11)
    doc.add_page_break()


def resumo(doc):
    h1(doc, "Resumo")
    para(doc,
         "Este relatório apresenta os resultados da organização, do "
         "pré-processamento e da análise, por métodos de aprendizado de máquina "
         "(machine learning, ML), de uma base de dados de qualidade da água bruta "
         "do rio Piracicaba, monitorada pelo SEMAE-Piracicaba entre 2009 e 2024. "
         "A base original, extraída dos boletins mensais em PDF, contém para cada "
         "mês três estatísticas (mínimo, médio e máximo) de 23 variáveis "
         "físico-químicas e bacteriológicas. A tabela de modelagem utilizou a "
         "estatística média mensal, resultando em 192 observações (16 anos × 12 "
         "meses). O fluxo compreendeu: (i) organização da base e análise "
         "exploratória inicial dos dados brutos — distribuições, sazonalidade e "
         "diagnóstico de dados faltantes —, que fundamentou as decisões de "
         "tratamento; (ii) pré-processamento propriamente dito, com o tratamento "
         "dos faltantes pela regra dos 30% (variáveis com mais de 30% de ausências "
         "são descartadas) e imputação das demais por média, mediana ou KNN — é "
         "nesta etapa que a imputação é realizada —, produzindo a tabela de "
         "modelagem completa de 192 meses, seguida da caracterização descritiva "
         "da base já limpa e imputada; (iii) estudo comparativo entre o método "
         "tradicional de remoção de outliers (IQR) e um ensemble de ML (Isolation "
         "Forest, LOF e KNN); e (iv) aplicação de três algoritmos de ML com "
         "objetivos complementares: K-Means, para identificar de forma não "
         "supervisionada regimes naturais de qualidade da água; Random Forest, "
         "para classificar os meses segundo a severidade de não conformidade com "
         "os limites da CONAMA e revelar as variáveis mais associadas a essa "
         "carga; e Isolation Forest, para detectar automaticamente meses anômalos "
         "(picos de contaminação), servindo de alerta para monitoramento "
         "contínuo. "
         "Os resultados mostram que a imputação preservou a estrutura dos dados, "
         "que o ensemble de ML é substancialmente mais preciso que os métodos "
         "univariados na sinalização de eventos críticos, e que o K-Means separa "
         "de forma consistente dois regimes de qualidade associados ao ciclo "
         "hidrológico (estiagem vs. chuvas). Todo o código foi implementado em "
         "Python, documentado e disponibilizado em repositório versionado, "
         "garantindo reprodutibilidade.")
    p = doc.add_paragraph()
    r = p.add_run("Palavras-chave: ")
    r.bold = True
    p.add_run("qualidade da água; irrigação; aprendizado de máquina; detecção de "
              "outliers; bacia PCJ; imputação de dados.")
    doc.add_page_break()


def introducao(doc):
    h1(doc, "1. Introdução")
    para(doc,
         "A agricultura é uma atividade econômica essencial e, ao mesmo tempo, de "
         "alto risco, exigindo estratégias de manejo que otimizem o uso dos "
         "recursos naturais. Diante da necessidade de aumento de produtividade e da "
         "maior variabilidade do regime de chuvas associada às mudanças climáticas, "
         "cresce a adoção da irrigação. O Brasil está entre os dez países com maior "
         "área equipada para irrigação (cerca de 8,2 milhões de hectares, "
         "responsáveis por 16% da produção agropecuária nacional), o que torna a "
         "gestão da qualidade da água de irrigação um tema estratégico, alinhado ao "
         "Objetivo de Desenvolvimento Sustentável n. 2 da ONU.")
    para(doc,
         "Não há consenso na literatura sobre quais variáveis melhor determinam a "
         "adequação da água à irrigação, mas destacam-se pH, condutividade elétrica, "
         "dureza, teor de ferro e manganês, cloretos, nutrientes e indicadores "
         "bacteriológicos. A má qualidade da água pode causar impermeabilização e "
         "degradação do solo, fitotoxicidade e redução do crescimento vegetal, além "
         "de entupimentos físicos, químicos e biológicos em sistemas de irrigação, "
         "elevando custos de manutenção. Monitorar esses parâmetros é, portanto, "
         "indispensável; porém os métodos tradicionais de avaliação são lentos e "
         "custosos.")
    para(doc,
         "A inteligência artificial, e em particular o aprendizado de máquina, pode "
         "acelerar esse processo, identificar padrões ocultos e apoiar decisões. "
         "Existe, contudo, uma lacuna de métodos automáticos para avaliar a "
         "qualidade da água para irrigação na região das bacias dos rios Piracicaba, "
         "Capivari e Jundiaí (PCJ). Este trabalho contribui para preencher essa "
         "lacuna organizando e analisando, com técnicas de ML, uma base histórica do "
         "SEMAE-Piracicaba, com foco em pré-processamento robusto, detecção de "
         "anomalias e extração de padrões.")
    h2(doc, "1.1 Objetivos")
    para(doc,
         "Objetivo geral: organizar e analisar, por métodos de ML, uma base de "
         "qualidade da água da bacia do PCJ (dados da ETA do rio Piracicaba, "
         "SEMAE, em parceria com o CENA).")
    para(doc, "Objetivos específicos:")
    for it in ("promover a limpeza e o pré-processamento da base;",
               "realizar análise exploratória e descritiva;",
               "comparar o método tradicional (IQR) e o de ML para remoção de "
               "outliers;",
               "desenvolver e aplicar modelos de ML para analisar as variáveis de "
               "qualidade da água e suas correlações."):
        doc.add_paragraph(it, style="List Bullet")


def metodologia(doc):
    h1(doc, "2. Materiais e Métodos")
    h2(doc, "2.1 Aquisição e organização dos dados a partir dos boletins em PDF")
    para(doc,
         "Os dados brutos foram disponibilizados pelo SEMAE-Piracicaba "
         "exclusivamente como boletins mensais em PDF, sem estruturação tabular, "
         "o que inviabilizava seu uso direto na análise. Desenvolveu-se, portanto, "
         "um script de extração (scripts/extrair_dados.py) baseado na biblioteca "
         "pdfplumber, responsável por ler as tabelas página a página e "
         "reorganizá-las no arquivo data/interim/dados_organizados.csv, seguindo o "
         "fluxo data/raw \u2192 data/interim adotado no projeto. Foram geradas 576 "
         "linhas (16 anos \u00d7 12 meses \u00d7 3 estatísticas \u2014 mínimo, médio "
         "e máximo), preservando a codificação UTF-8 dos rótulos (Mín., Méd., Máx., "
         "Cianobactéria). Das 23 variáveis do schema, 22 foram mapeadas para o "
         "período 2009\u20132020 \u2014 a Amônia está legitimamente ausente nesses "
         "anos \u2014 e as 23 para 2021\u20132024.")
    para(doc,
         "A principal dificuldade foi a heterogeneidade do layout ao longo dos 16 "
         "anos. O script trata esses casos por meio de um mapa de cabeçalhos "
         "(HEADER_MAP) com normalização de texto (remoção de espaços, parênteses e "
         "quebras de linha e correção do encoding latin-1 corrompido do PDF), o que "
         "permitiu: (a) reconciliar o layout reordenado de 2011, ignorando a coluna "
         "extra \u201cFenol\u201d e reposicionando \u201cCond. (\u00b5S/cm)\u201d e "
         "\u201cCianobactérias\u201d; (b) tratar a renomeação de \u201cC.F.\u201d "
         "para \u201cE. coli\u201d em 2018\u20132020, mapeando-a para a mesma "
         "variável; e (c) acomodar a inclusão da Amônia a partir de 2021 como nova "
         "coluna do schema, deixada vazia nos anos anteriores. O mês de referência "
         "passou a ser lido da coluna inicial da linha de mínimo (e não da linha de "
         "média, como assumia a versão preliminar). O script expõe uma interface de "
         "linha de comando (--pdf, --out, -v) com caminhos padrão relativos à raiz "
         "do projeto, garantindo reprodutibilidade.")
    h2(doc, "2.2 Carregamento, ambiente computacional e critérios metodológicos")
    para(doc,
         "A base organizada (data/interim/dados_organizados.csv) foi carregada com "
         "um parser numérico brasileiro que trata separador decimal por vírgula, "
         "separador de milhar por ponto e marcadores de ausência (--, ---, -). A "
         "tabela de modelagem usa a estatística média mensal, resultando em 192 "
         "observações com data e índices temporais (mês, trimestre e estação "
         "seca/úmida).")
    para(doc,
         "Todo o processamento foi feito em Python, com código "
         "modular em src/ e scripts de pipeline em scripts/. Os principais pacotes "
         "são pandas, numpy, scikit-learn, scipy, matplotlib/seaborn e pyod. A "
         "semente aleatória foi fixada (random_state = 42) para reprodutibilidade; "
         "figuras são salvas em reports/figures/ e tabelas em reports/tables/. Os "
         "códigos estão documentados em notebooks Jupyter e versionados em "
         "repositório Git.")
    para(doc, "As regras metodológicas seguem o projeto.")
    para(doc,
         "(a) Tratamento de dados faltantes, conforme a proporção de ausências em "
         "cada variável:")
    _df_falt = pd.DataFrame({
        "cond": ["Acima de 30%", "Entre 5% e 30%", "Abaixo de 5%"],
        "aplic": [
            "A variável é descartada (há poucos dados confiáveis para recuperar).",
            "O valor ausente é estimado por KNN: preenchido a partir dos k = 5 "
            "meses mais parecidos (os \u201cvizinhos mais próximos\u201d), com "
            "mais peso aos mais semelhantes; as variáveis são padronizadas quando "
            "os faltantes passam de 15%.",
            "O valor ausente é preenchido pela média, quando a distribuição é "
            "simétrica, ou pela mediana, quando é assimétrica (|skew| > 1, com "
            "cauda longa para um dos lados), por ser menos sensível a valores "
            "extremos.",
        ],
    })
    tabela(doc, _df_falt, "", col_labels=["Condição", "Aplicação"], num=False)
    para(doc, "(b) Detecção de outliers, por etapas:")
    _df_out = pd.DataFrame({
        "etapa": ["Algoritmos base", "Combinação (ensemble)", "Calibração"],
        "desc": [
            "Vários detectores independentes (IQR, MAD, Isolation Forest, LOF e "
            "KNN) analisam cada mês e apontam os pontos que consideram anômalos.",
            "Os detectores formam um \u201ccomitê\u201d: um mês só é marcado como "
            "outlier quando a maioria concorda (votação).",
            "A fração esperada de anomalias (contamination) é ajustada por "
            "validação cruzada \u2014 testando o modelo em diferentes subconjuntos "
            "dos dados para escolher o valor mais estável.",
        ],
    })
    tabela(doc, _df_out, "", col_labels=["Etapa", "Descrição"], num=False)
    para(doc, "(c) Validação dos resultados:")
    _df_val = pd.DataFrame({
        "alvo": ["Imputação", "Agrupamento (K-Means)",
                 "Classificação (Random Forest)"],
        "como": [
            "Comparação da média e do desvio antes e depois do preenchimento, para "
            "confirmar que a imputação não distorceu os dados.",
            "Coeficiente de silhueta (quão coesos e separados ficam os grupos, de "
            "-1 a 1) e soma das distâncias internas de cada grupo (WCSS \u2014 "
            "quanto menor, mais compactos).",
            "Acurácia (percentual de acertos) e F1 (equilíbrio entre precisão e "
            "sensibilidade, útil quando as classes são desbalanceadas).",
        ],
    })
    tabela(doc, _df_val, "", col_labels=["Etapa avaliada", "Como é validada"], num=False)
    para(doc,
         "Os limites regulatórios de referência são os da Resolução CONAMA "
         "357/2005, Classe 2.")
    doc.add_page_break()


def sec_preproc(doc):
    h1(doc, "3. Resultados e Discussão")
    h2(doc, "3.1 Pré-processamento e dados faltantes")
    para(doc,
         "A distribuição de faltantes é fortemente desigual. Sete variáveis "
         "apresentam mais de 30% de ausências — sólidos totais e surfactantes são "
         "quase nunca medidos (~99%), e nutrientes (P, N) e amônia ficam acima de "
         "75% —, enquanto os parâmetros físico-químicos básicos (pH, turbidez, cor) "
         "têm menos de 1%. As Figuras 1 a 3 evidenciam que a ausência é "
         "estrutural por período: várias variáveis só passam a ser medidas de forma "
         "regular a partir de 2018–2021, refletindo mudanças de protocolo do "
         "laboratório, e não falhas aleatórias.")
    figura(doc, "missing_pct_barras.png",
           "Percentual de valores ausentes por variável (192 meses).")
    figura(doc, "missing_matriz.png",
           "Matriz de ausência (linhas = meses, colunas = variáveis): o padrão em "
           "blocos indica ausência estrutural por período.")
    figura(doc, "missing_temporal.png",
           "Cobertura temporal: momento em que cada variável passou a ser medida.")
    para(doc,
         "O projeto assume o mecanismo MCAR (Missing Completely At Random). O teste "
         "de Little resultou em qui-quadrado ~ 327,6 (gl = 269, p ~ 0,008). Como "
         "p < 0,05, a hipótese MCAR é rejeitada: os faltantes não são completamente "
         "aleatórios, o que é coerente com o padrão temporal observado. Essa "
         "limitação fica registrada para a interpretação dos resultados; ainda "
         "assim, aplicam-se as regras de imputação declaradas, pois a maioria das "
         "variáveis mantidas tem baixíssimo percentual de ausência.")
    df = _read("faltantes_resumo.csv")
    if df is not None:
        tabela(doc, df, "Decisão de imputação por variável (ordenada por % de "
               "faltantes).",
               col_labels=["Variável", "% NaN", "Assimetria", "Decisão"])
    para(doc,
         "A aplicação da regra resulta em 7 variáveis excluídas (sólidos totais, "
         "surfactantes, DBO, P, N, amônia e coliformes totais) e 16 mantidas para a "
         "modelagem. Registra-se o conflito de que a regra descarta justamente "
         "nutrientes (N, P) e amônia, relevantes para irrigação; uma mitigação "
         "possível seria um pipeline alternativo restrito ao período 2018–2024, "
         "quando essas variáveis passam a ser medidas com regularidade.")


def sec_imputacao(doc):
    h2(doc, "3.2 Imputação e sua validação")
    para(doc,
         "As variáveis com menos de 5% de faltantes foram imputadas por média "
         "(pH e OD, simétricas) ou mediana (as demais, assimétricas). As quatro "
         "variáveis na faixa 5–30% (coliformes fecais, fluoreto, clorofila e "
         "cianobactérias) foram imputadas por KNN (k = 5, pesos por distância). A "
         "Tabela 2 compara média, mediana e desvio antes e depois da imputação: as "
         "variações são pequenas (em geral abaixo de 10% no desvio), indicando que "
         "a imputação não distorceu as distribuições.")
    df = _read("validacao_imputacao.csv")
    if df is not None:
        df = df[["variavel", "media_antes", "media_depois", "std_antes",
                 "std_depois", "delta_media_%", "delta_std_%"]]
        tabela(doc, df, "Estatísticas antes e depois da imputação (16 variáveis "
               "mantidas).",
               col_labels=["Variável", "Média antes", "Média depois",
                           "Desvio antes", "Desvio depois", "Δ média %",
                           "Δ desvio %"])
    para(doc,
         "Além da comparação de momentos, avaliou-se o impacto da imputação na "
         "estrutura de agrupamento (Tabela 3). A silhueta do conjunto imputado "
         "(0,279, com 192 observações) é superior à do subconjunto complete-case "
         "(0,203, com 126 observações). Ou seja, recuperar as linhas com faltantes "
         "não apenas preservou como reforçou uma separação consistente dos dados, "
         "validando a estratégia de imputação.")
    df = _read("validacao_imputacao_kmeans.csv")
    if df is not None:
        tabela(doc, df, "Impacto da imputação na silhueta do K-Means.",
               col_labels=["Conjunto", "n", "Melhor k", "Silhueta"])
    para(doc,
         "Por fim, os dados foram padronizados (StandardScaler, média 0 e desvio 1) "
         "para os algoritmos sensíveis à escala (K-Means, KNN e LOF). A Figura 4 "
         "apresenta a matriz de correlação de Pearson entre as 16 variáveis "
         "mantidas, com as associações esperadas: dureza × condutividade × "
         "alcalinidade (mineralização conjunta) e cor × turbidez (material "
         "particulado).")
    figura(doc, "eval_correlacao.png",
           "Matriz de correlação de Pearson entre as 16 variáveis mantidas.")
    para(doc,
         "Como a base não possui referência externa, a qualidade da imputação foi "
         "validada diretamente por mascaramento (hold-out): esconde-se de forma "
         "aleatória (MCAR) uma fração de 15% das células efetivamente medidas, "
         "imputa-se por cada método e mede-se o erro apenas nessas células — cujo "
         "valor verdadeiro é conhecido. O procedimento é repetido 30 vezes "
         "(Monte Carlo) com as mesmas máscaras para todos os métodos, resumindo-se "
         "o erro por RMSE, MAE e nRMSE (RMSE ÷ desvio da variável, comparável entre "
         "escalas). A ordem metodológica é preservada: primeiro identifica-se o "
         "percentual de faltantes por variável (Tabela 1) e a faixa da regra; só "
         "então a imputação é aplicada e validada. O método vencedor por "
         "variável consta na Tabela 4 e o erro por método, na Figura 5.")
    df = _read("imputacao_melhor_metodo.csv")
    if df is not None:
        df = df[["variavel", "melhor_nrmse", "melhor_mae", "metodo_regra"]]
        tabela(doc, df, "Melhor método de imputação por variável no teste de "
               "mascaramento (por nRMSE e por MAE) vs. método aplicado pela regra.",
               col_labels=["Variável", "Melhor (nRMSE)", "Melhor (MAE)",
                           "Método da regra"])
    figura(doc, "imputacao_benchmark_nrmse.png",
           "Erro de imputação (nRMSE) por variável e método no teste de "
           "mascaramento; menor é melhor.", largura=6.2)
    para(doc,
         "Três leituras emergem. Primeiro, o KNN — por ser multivariado e explorar "
         "as correlações entre variáveis — costuma ter o menor erro de predição, "
         "mesmo em variáveis com poucos faltantes; ainda assim a regra reserva o "
         "KNN à faixa de 5–30% por parcimônia, pois abaixo de 5% pouquíssimas "
         "células são imputadas e uma solução univariada robusta é suficiente. "
         "Segundo, entre média e mediana o método vencedor depende da métrica: o "
         "RMSE favorece a média (que minimiza o erro quadrático por construção), "
         "enquanto o MAE — robusto à assimetria — favorece a mediana justamente nas "
         "variáveis assimétricas em que a regra a seleciona, o que confirma a "
         "escolha do projeto. Terceiro, os níveis de nRMSE (~0,5–1,0) mostram que o "
         "erro por célula não é desprezível, algo típico de dados ambientais; "
         "porém, como apenas uma pequena fração das células das variáveis mantidas "
         "é de fato imputada, o impacto no conjunto é reduzido — coerente com as "
         "variações mínimas de média e desvio da Tabela 2 e com a silhueta "
         "preservada da Tabela 3.")


def sec_eda(doc):
    h2(doc, "3.3 Análise descritiva e exploratória")
    para(doc,
         "A análise descritiva resume a informação das 16 variáveis mantidas "
         "(Tabela 5). Chamam a atenção a forte assimetria positiva de várias "
         "variáveis — cor, cloreto, coliformes fecais, clorofila e cianobactérias "
         "apresentam |skew| elevado —, o que justifica o uso da mediana na "
         "imputação e a atenção a outliers. O pH permanece em faixa estreita "
         "(média 7,3; desvio 0,18), enquanto o ferro (média 2,1 mg/L) excede "
         "sistematicamente o limite CONAMA de referência, sinalizando forte "
         "presença de material ferruginoso na água bruta.")
    df = _read("descritiva_processada.csv")
    if df is not None:
        df = df[["variavel", "media", "desvio", "minimo", "mediana", "maximo",
                 "assimetria"]]
        tabela(doc, df, "Estatística descritiva das 16 variáveis mantidas "
               "(base imputada).",
               col_labels=["Variável", "Média", "Desvio", "Mín.", "Mediana",
                           "Máx.", "Assimetria"])
    para(doc,
         "Os histogramas (Figura 6) confirmam distribuições assimétricas e de "
         "cauda longa na maioria das variáveis. Os boxplots padronizados (Figura 7) "
         "permitem comparar as variáveis em escala comum e revelam numerosos pontos "
         "extremos, especialmente em cor, turbidez, cloreto e nos indicadores "
         "biológicos. As séries temporais (Figura 8) mostram a dinâmica mensal de "
         "variáveis-chave, com picos de turbidez e de condutividade coerentes com "
         "eventos hidrológicos (cheias e estiagens).")
    figura(doc, "eda_histogramas.png",
           "Histogramas das 16 variáveis mantidas (base imputada).", largura=6.3)
    figura(doc, "eda_boxplots.png",
           "Boxplots padronizados (z-score) das 16 variáveis: comparação de "
           "dispersão e pontos extremos.", largura=6.3)
    figura(doc, "eda_series_temporais.png",
           "Séries temporais mensais de variáveis-chave (2009–2024).", largura=6.0)


def sec_outliers(doc):
    h2(doc, "3.4 Detecção de outliers: método tradicional vs. ensemble de ML")
    para(doc,
         "Um dos objetivos específicos é comparar o método tradicional de remoção "
         "de outliers com uma abordagem de ML. Confrontaram-se dois métodos "
         "univariados (IQR, de Tukey, e o desvio absoluto mediano, MAD) com "
         "detectores multivariados (Isolation Forest, LOF e KNN) e o ensemble por "
         "votação. A Tabela 6 resume o desempenho. Os métodos univariados "
         "sinalizam quase metade das observações como outliers (IQR 47,6%; MAD "
         "50,3%), o que é excessivo e pouco útil operacionalmente; já os detectores "
         "de ML sinalizam de 5% a 7% e concentram-se nos pontos mais anômalos.")
    df = pd.read_csv(PROJECT_ROOT / "data" / "interim" / "comparacao_metodos.csv")
    tabela(doc, df, "Comparação entre métodos de detecção de outliers (192 "
           "observações).",
           col_labels=["Método", "N flags", "Taxa", "Dist. calib.",
                       "Recall crise", "Precisão crise", "Concord. média"])
    para(doc,
         "Tomando o período da crise hídrica de 2014–2016 como referência de "
         "eventos críticos, os métodos univariados têm maior recall (0,64–0,67) mas "
         "baixíssima precisão (~0,25) — marcam muitos pontos, acertando os "
         "críticos por força do volume. Os métodos de ML invertem esse "
         "compromisso: maior precisão (LOF 0,54; IForest 0,62) e o ensemble atinge "
         "a melhor precisão de crise (0,67), ao custo de menor recall. Para "
         "sinalizar eventos que justifiquem inspeção, a maior precisão do ensemble "
         "é preferível. A Figura 9 mostra a contagem por método e a Figura 10, o "
         "diagrama de Venn da sobreposição entre IQR, MAD e ensemble.")
    figura(doc, "metodos_contagem.png",
           "Número de outliers sinalizados por cada método.", largura=5.6)
    figura(doc, "metodos_venn.png",
           "Sobreposição dos conjuntos de outliers (IQR, MAD e ensemble).",
           largura=4.8)
    para(doc,
         "O parâmetro contamination do ensemble foi calibrado por validação "
         "cruzada, medindo a estabilidade (índice de Jaccard) das flags entre os "
         "folds (Tabela 7 e Figura 11). O valor mais estável foi 0,08 (Jaccard "
         "0,833), coerente com a faixa 0,05–0,10 do projeto. Cruzando as camadas "
         "univariada e multivariada, obtêm-se 9 outliers confirmados — pontos "
         "sinalizados de forma consistente por múltiplos métodos —, cuja "
         "distribuição temporal (Figura 12) concentra-se em eventos de 2009 e no "
         "período crítico de 2014–2016.")
    df = _read("contamination_calibracao.csv")
    if df is not None:
        tabela(doc, df, "Calibração do parâmetro contamination por estabilidade "
               "(validação cruzada).",
               col_labels=["Contamination", "N flags", "Taxa",
                           "Estabilidade (Jaccard)"])
    figura(doc, "contamination_calibracao.png",
           "Estabilidade (Jaccard) das flags em função do contamination.",
           largura=5.6)
    figura(doc, "outliers_confirmados_timeline.png",
           "Linha do tempo dos outliers confirmados pelo cruzamento das camadas.",
           largura=6.0)


def sec_modelos(doc):
    h2(doc, "3.5 Modelagem: K-Means, Random Forest e Isolation Forest")
    para(doc,
         "K-Means (agrupamento). A varredura de k (Figura 13) indica k = 2 pela "
         "silhueta (0,279) e k = 4 pelo método do cotovelo; adotou-se k = 2, o mais "
         "estável. Embora a silhueta esteja abaixo do limiar de 0,5 (agrupamento "
         "moderado), os dois grupos são interpretáveis (Tabela 8). O Cluster 0 "
         "(n = 33) reúne meses de água mais mineralizada — condutividade, "
         "alcalinidade, cloreto e dureza elevados, oxigênio dissolvido baixo e "
         "maior carga de clorofila/cianobactérias —, típico de regime de estiagem "
         "e baixa vazão. O Cluster 1 (n = 158) reúne meses de maior turbidez e cor "
         "com menor mineralização, associados à diluição por chuvas. A Figura 14 "
         "ilustra a separação dos clusters.")
    df = _read("kmeans_perfil.csv")
    if df is not None:
        dft = df.set_index("cluster").T.reset_index()
        dft.columns = ["variavel"] + [f"cluster_{c}" for c in df["cluster"].tolist()]
        tabela(doc, dft, "Perfil médio de cada cluster (K-Means, k = 2).",
               col_labels=["Variável"] + [f"Cluster {c}" for c in
                                          df["cluster"].tolist()], fontsize=8)
    figura(doc, "eval_kmeans_cotovelo_silhueta.png",
           "Método do cotovelo (WCSS) e coeficiente de silhueta em função de k.",
           largura=6.2)
    figura(doc, "eval_kmeans_scatter.png",
           "Dispersão das observações coloridas por cluster.", largura=5.4)
    para(doc,
         "Random Forest (classificação). Como a água é bruta, o rótulo estrito de "
         "adequação pela CONAMA é degenerado (todas as 192 amostras violam algum "
         "limite). Adotou-se, então, um alvo de severidade: número de parâmetros "
         "violados binarizado pela mediana (alta vs. baixa carga de não "
         "conformidade). Para evitar vazamento de dados, as variáveis usadas na "
         "regra de rotulagem foram excluídas das features, restando 5 preditoras "
         "(alcalinidade, acidez, oxigênio consumido, dureza e condutividade). O "
         "modelo obteve acurácia 0,692 e F1 0,400 (CV-F1 média 0,434) — Tabela 9. "
         "O desempenho moderado é esperado dado o pequeno número de amostras e de "
         "features não vazadas; ainda assim, as importâncias (Figura 15) destacam "
         "alcalinidade e condutividade como preditoras dominantes da severidade, "
         "coerente com a interpretação mineral dos clusters.")
    df = _read("rf_metricas.csv")
    if df is not None:
        tabela(doc, df, "Métricas do Random Forest (alvo de severidade CONAMA).",
               col_labels=["Acurácia", "Precisão", "Recall", "F1"])
    df = _read("rf_importancias.csv")
    if df is not None:
        tabela(doc, df, "Importância das variáveis no Random Forest.",
               col_labels=["Variável", "Importância"])
    figura(doc, "eval_rf_importancias.png",
           "Importância das variáveis (Random Forest).", largura=5.8)
    figura(doc, "eval_rf_matriz_confusao.png",
           "Matriz de confusão do Random Forest (conjunto de teste).",
           largura=4.6)
    para(doc,
         "Isolation Forest (anomalias). Aplicado sobre os dados padronizados com "
         "contamination 0,07, o modelo sinalizou 14 anomalias entre 192 meses "
         "(taxa 0,073) — Figura 17. As anomalias coincidem em grande parte com os "
         "outliers confirmados na seção anterior, reforçando a consistência entre "
         "as abordagens e a utilidade do método para monitoramento contínuo e "
         "alerta de picos de contaminação.")
    figura(doc, "eval_iforest_scatter.png",
           "Anomalias sinalizadas pelo Isolation Forest (destaque em vermelho).",
           largura=5.4)


def conclusao(doc):
    h1(doc, "4. Conclusões")
    para(doc,
         "O trabalho organizou e analisou, de ponta a ponta e de forma "
         "reprodutível, uma base histórica de qualidade da água bruta do rio "
         "Piracicaba. No pré-processamento, a regra dos 30% e a imputação "
         "combinada (média/mediana e KNN) mantiveram 16 das 23 variáveis sem "
         "distorcer as distribuições; a validação pela silhueta mostrou que a "
         "imputação até reforçou a estrutura dos dados. O teste de Little indicou "
         "que a ausência não é completamente aleatória (mecanismo ligado a "
         "mudanças de protocolo), limitação devidamente registrada.")
    para(doc,
         "No estudo comparativo de outliers, o ensemble de ML mostrou-se "
         "claramente superior aos métodos univariados em precisão para sinalizar "
         "eventos críticos, enquanto IQR e MAD superestimam grosseiramente a "
         "quantidade de outliers. A calibração por validação cruzada fixou "
         "contamination em 0,08, e o cruzamento de camadas produziu um conjunto "
         "enxuto de 9 outliers confirmados, concentrados na crise hídrica de "
         "2014–2016.")
    para(doc,
         "Na modelagem, o K-Means separou dois regimes de qualidade coerentes com "
         "o ciclo hidrológico (estiagem mineralizada vs. chuvas com diluição e "
         "turbidez), o Random Forest identificou alcalinidade e condutividade como "
         "preditoras dominantes da severidade de não conformidade, e o Isolation "
         "Forest sinalizou anomalias consistentes com os eventos críticos. Em "
         "conjunto, os resultados demonstram que técnicas de ML — especialmente as "
         "não supervisionadas — agregam valor à avaliação da qualidade da água e "
         "podem apoiar a tomada de decisão.")
    para(doc,
         "Como limitações, destacam-se o número reduzido de observações mensais, a "
         "exclusão de nutrientes relevantes para irrigação por falta de "
         "granularidade histórica, e o fato de os limites CONAMA se referirem ao "
         "corpo d'água e não à água já tratada. Trabalhos futuros incluem um "
         "pipeline restrito a 2018–2024 (com nutrientes), a incorporação de índices "
         "consagrados de qualidade de água para irrigação (por exemplo, SAR e RAS) "
         "e a integração com sensoriamento e IoT para monitoramento em tempo real.")


def referencias(doc):
    h1(doc, "5. Referências (seleção)")
    refs = [
        "AGGARWAL, C. C. Outlier Analysis. 2. ed. Springer, 2017.",
        "AYERS, R. S.; WESTCOT, D. W. Water quality for agriculture. FAO Irrigation "
        "and Drainage Paper 29, 1985.",
        "CONAMA. Resolução n. 357, de 17 de março de 2005. Ministério do Meio "
        "Ambiente, Brasil.",
        "LITTLE, R. J. A. A test of missing completely at random for multivariate "
        "data with missing values. JASA, 83(404), 1988.",
        "LIU, F. T.; TING, K. M.; ZHOU, Z.-H. Isolation Forest. ICDM, 2008.",
        "MORGENTHALER, S. Exploratory data analysis. WIREs Comp. Stat., 2009.",
        "ROUSSEEUW, P. J. Silhouettes: a graphical aid to the interpretation and "
        "validation of cluster analysis. J. Comput. Appl. Math., 1987.",
        "TROYANSKAYA, O. et al. Missing value estimation methods for DNA "
        "microarrays. Bioinformatics, 2001.",
        "VAN BUUREN, S. Flexible Imputation of Missing Data. 2. ed. CRC Press, 2018.",
        "ZHAO, Y.; NASRULLAH, Z.; LI, Z. PyOD: A Python toolbox for scalable "
        "outlier detection. JMLR, 2019/2020.",
    ]
    for r in refs:
        p = doc.add_paragraph(r)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def cronograma(doc):
    h1(doc, "6. Cronograma de execução")
    dados = [
        ("1. Revisão bibliográfica", "Concluída"),
        ("2. Pré-processamento", "Concluída"),
        ("3. Imputação e retirada de outliers", "Concluída"),
        ("4. Análise descritiva", "Concluída"),
        ("5. Aplicação de modelos", "Concluída"),
        ("6. Análise das métricas", "Concluída"),
        ("7. Relatório final", "Concluída"),
    ]
    df = pd.DataFrame(dados, columns=["Atividade", "Situação"])
    tabela(doc, df, "Situação das atividades do cronograma.",
           col_labels=["Atividade", "Situação"], fontsize=10)


def rodape_paginas(doc):
    sec = doc.sections[0]
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    instr = run._r.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
    instr.text = "PAGE"
    fldChar2 = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run._r.append(fldChar1)
    run._r.append(instr)
    run._r.append(fldChar2)


def main() -> None:
    logging.basicConfig(level=logging.INFO, format="%(levelname)s | %(message)s")
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    capa(doc)
    resumo(doc)
    introducao(doc)
    metodologia(doc)
    sec_preproc(doc)
    sec_imputacao(doc)
    sec_eda(doc)
    sec_outliers(doc)
    sec_modelos(doc)
    conclusao(doc)
    referencias(doc)
    cronograma(doc)
    rodape_paginas(doc)

    saida = config.REPORTS_DIR / "relatorio_final.docx"
    doc.save(str(saida))
    LOG.info("Relatorio salvo: %s (%d figuras, %d tabelas)", saida, _fig_n, _tab_n)


if __name__ == "__main__":
    main()
